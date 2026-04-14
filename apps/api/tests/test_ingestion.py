"""Ingestion service — focus on TXT + DOCX paths that don't require poppler."""

from __future__ import annotations

from pathlib import Path

import pytest

from app.services.ingestion import (
    MIN_EXTRACTED_CHARS,
    extract_text,
    find_page_for_offset,
)

SAMPLE = (
    "FREELANCE SERVICE AGREEMENT\n"
    "1. PAYMENT: Client will pay upon final delivery. Client may withhold\n"
    "   payment if work does not meet expectations, at Client's sole discretion.\n"
    "2. TERMINATION: Client may terminate this agreement at any time without cause.\n"
)


def test_extract_txt(tmp_path: Path) -> None:
    f = tmp_path / "contract.txt"
    f.write_text(SAMPLE, encoding="utf-8")
    doc = extract_text(f)
    assert doc.source_format == "txt"
    assert "FREELANCE SERVICE AGREEMENT" in doc.text
    assert doc.page_map[0]["page"] == 1
    assert doc.ocr_used is False
    assert doc.preview.startswith("FREELANCE")


def test_extract_missing_file(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        extract_text(tmp_path / "nope.txt")


def test_unsupported_extension(tmp_path: Path) -> None:
    f = tmp_path / "mystery.xyz"
    f.write_text("hi")
    with pytest.raises(ValueError):
        extract_text(f)


def test_find_page_for_offset() -> None:
    page_map = [
        {"page": 1, "start": 0, "end": 100},
        {"page": 2, "start": 100, "end": 250},
    ]
    assert find_page_for_offset(page_map, 50) == 1
    assert find_page_for_offset(page_map, 150) == 2
    assert find_page_for_offset(page_map, 999) is None


def test_extract_docx(tmp_path: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    f = tmp_path / "contract.docx"
    doc = Document()
    doc.add_paragraph("FREELANCE SERVICE AGREEMENT")
    doc.add_paragraph("1. PAYMENT: Client will pay upon final delivery.")
    doc.save(str(f))

    result = extract_text(f)
    assert result.source_format == "docx"
    assert "FREELANCE SERVICE AGREEMENT" in result.text
    assert len(result.text) >= MIN_EXTRACTED_CHARS - 50  # short doc, still OK
