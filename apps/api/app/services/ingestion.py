"""
File ingestion — turn uploaded PDF/DOCX/TXT into plain text + page map.

Strategy:
- PDF:  pdfplumber (fast, accurate for digital PDFs). If combined extracted text
        is shorter than MIN_EXTRACTED_CHARS, fall back to pytesseract OCR.
- DOCX: python-docx (paragraphs).
- TXT:  read utf-8, ignore errors.

Returns (full_text, page_map) where page_map is a list of
{"page": int, "start": int, "end": int} describing character ranges of each
page inside `full_text`. For non-paginated formats (DOCX, TXT) we use a single
page=1 entry.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, TypedDict

logger = logging.getLogger(__name__)

MIN_EXTRACTED_CHARS = 100


class PageRange(TypedDict):
    page: int
    start: int
    end: int


@dataclass
class ExtractedDocument:
    """Result of ingestion."""

    text: str
    page_map: list[PageRange]
    source_format: str  # "pdf" | "docx" | "txt"
    ocr_used: bool = False

    @property
    def preview(self) -> str:
        """First ~500 chars for UI preview."""
        return self.text[:500]


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def extract_text(path: str | Path) -> ExtractedDocument:
    """Dispatch to the right extractor based on file extension."""

    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(p)
    if ext in {".docx", ".doc"}:
        return _extract_docx(p)
    if ext in {".txt", ".md", ""}:
        return _extract_txt(p)
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> ExtractedDocument:
    import pdfplumber

    parts: list[str] = []
    page_map: list[PageRange] = []
    cursor = 0

    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw = page.extract_text() or ""
            parts.append(raw)
            start = cursor
            cursor += len(raw) + 2  # +2 for the "\n\n" joiner we'll add later
            page_map.append({"page": i, "start": start, "end": start + len(raw)})

    text = "\n\n".join(parts).strip()

    if len(text) >= MIN_EXTRACTED_CHARS:
        return ExtractedDocument(text=text, page_map=page_map, source_format="pdf")

    # Scanned PDF — OCR fallback
    logger.info("PDF text too short (%d chars) — falling back to OCR", len(text))
    ocr_text, ocr_map = _ocr_pdf(path)
    return ExtractedDocument(
        text=ocr_text,
        page_map=ocr_map,
        source_format="pdf",
        ocr_used=True,
    )


def _ocr_pdf(path: Path) -> tuple[str, list[PageRange]]:
    """Convert each page to an image and OCR it."""

    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError as e:
        raise RuntimeError("OCR requires pytesseract and pdf2image installed.") from e

    images = convert_from_path(str(path))
    parts: list[str] = []
    page_map: list[PageRange] = []
    cursor = 0

    for i, image in enumerate(images, start=1):
        raw = pytesseract.image_to_string(image) or ""
        parts.append(raw)
        start = cursor
        cursor += len(raw) + 2
        page_map.append({"page": i, "start": start, "end": start + len(raw)})

    return "\n\n".join(parts).strip(), page_map


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs).strip()

    return ExtractedDocument(
        text=text,
        page_map=[{"page": 1, "start": 0, "end": len(text)}],
        source_format="docx",
    )


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def _extract_txt(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return ExtractedDocument(
        text=text,
        page_map=[{"page": 1, "start": 0, "end": len(text)}],
        source_format="txt",
    )


def find_page_for_offset(page_map: list[PageRange], offset: int) -> Optional[int]:
    """Given a character offset inside extracted text, return the 1-indexed page."""

    for entry in page_map:
        if entry["start"] <= offset < entry["end"]:
            return entry["page"]
    return None
